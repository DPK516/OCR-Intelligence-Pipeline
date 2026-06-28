import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK 
from docx.shared import Pt

def create_word_document(pages_structured_data, output_filename):
    print("\nStarting Advanced Layout-Aware Document Assembly...")
    doc = Document()
    
    for index, page_data in enumerate(pages_structured_data):
        page_number = index + 1
        print(f"Assembling Page {page_number} with advanced formatting...")
        
        for block in page_data.get("page_layout_corrected", []):
            block_type = block.get("type", "paragraph")
            
            
            if block_type == "table":
                rows_data = block.get("rows", [])
                if rows_data:
                    
                    num_rows = len(rows_data)
                    num_cols = len(rows_data[0]) if num_rows > 0 else 1
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Table Grid' 
                    table.autofit = True
                    
                    
                    for row_idx, row_text_list in enumerate(rows_data):
                        for col_idx, cell_text in enumerate(row_text_list):
                            
                            if col_idx < num_cols: 
                                clean_text = str(cell_text)
                                
                                if clean_text == "[BLANK]":
                                    clean_text = ""
                                table.cell(row_idx, col_idx).text = clean_text
                
                
                doc.add_paragraph()
                continue 
            
            
            p = doc.add_paragraph()
            
            
            alignment = block.get("alignment", "left")
            if alignment == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment == "right":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
            
            text_spans = block.get("text_spans", [])
            for span in text_spans:
                run = p.add_run(span.get("text", ""))
                
                
                if span.get("bold", False):
                    run.bold = True
                if span.get("italic", False):
                    run.italic = True
                    
                 
                if block.get("font_style") == "serif":
                    run.font.name = 'Times New Roman'
                else:
                    run.font.name = 'Arial'
        
        
        if page_number < len(pages_structured_data):
            if 'p' in locals():
                p.add_run().add_break(WD_BREAK.PAGE)
            else:
                doc.add_page_break()
            
    output_dir = "data/output_docs"
    os.makedirs(output_dir, exist_ok=True)
    final_path = os.path.join(output_dir, output_filename)
    
    try:
        doc.save(final_path)
        print(f"\nSUCCESS! Visually matched document saved at: {final_path}")
    except PermissionError:
        print(f"\n ERROR: Cannot save. Please close '{output_filename}' in Microsoft Word.")